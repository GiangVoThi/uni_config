from app.models.database import db
from app.models.contract_model import dichVu, noiDungCongViec, quyen, nghiaVu, phuongThucThanhToan, phiDichVu
import os
from docx import Document
from config import Config
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sqlalchemy import text
from datetime import datetime

def get_service_details(service_ids):
    # Chuyển list service_ids thành string để dùng trong câu SQL
    service_ids_str = ','.join(str(id) for id in service_ids)
    
    # Câu truy vấn lấy thông tin dịch vụ
    service_query = text(f"""
        SELECT id_dich_vu, ten_dich_vu 
        FROM dich_vu 
        WHERE id_dich_vu IN ({service_ids_str})
    """)
    
    details = {}
    services = db.session.execute(service_query)
    
    for service in services:
        # Câu truy vấn lấy nội dung công việc
        work_content_query = text("""
            SELECT ten_noi_dung, chi_tiet 
            FROM noi_dung_cong_viec 
            WHERE id_dich_vu = :service_id
        """)
        work_contents = db.session.execute(work_content_query, {'service_id': service.id_dich_vu})
        
        
        details[service.id_dich_vu] = {
            'ten_dich_vu': service.ten_dich_vu,
            'noi_dung': [(wc.ten_noi_dung, wc.chi_tiet) for wc in work_contents]
        }
    
    return details

def preview_contract(service_details, contract_number, sign_date):
    preview_data = {
        'header': {
            'hanoi': 'Hanoi: 11 Dich Vong Hau, Cau Giay District',
            'hcm': 'Ho Chi Minh: 105D Ben Van Don, District 4',
            'phone': 'Phone: +(84) 908-535-898',
            'email': 'Email: univn@exumuni.com'
        },
        'title': {
            'vn': 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM',
            'vn_sub': 'Độc lập - Tự do - Hạnh phúc',
            'en': 'SOCIALIST REPUBLIC OF VIETNAM',
            'en_sub': 'Independence - Freedom - Happiness'
        },
        'contract': {
            'vn': 'HỢP ĐỒNG DỊCH VỤ',
            'vn_number': f'Số: {contract_number}',
            'en': 'SERVICE CONTRACT',
            'en_number': f'No. {contract_number}'
        },
        'date': {
            'vn': f'Hợp đồng này được lập ngày {sign_date} tại trụ sở của CÔNG TY TNHH TƯ VẤN UNI.\nSau khi bàn bạc thỏa thuận, hai bên nhất trí ký Hợp đồng dịch vụ tư vấn với các điều khoản sau:',
        },
        'intro': {
            'en': f'This contract was made on {sign_date} at UNI CONSULTING CO., LTD.\nAfter discussing the agreement, the two sides agreed to sign a consulting service contract with the following terms:'
        },
        'dieu_3': []  # Nội dung công việc
    }
    
    # Xử lý Điều 3: Nội dung công việc
    num_services = len(service_details)
    for idx, (service_id, data) in enumerate(service_details.items(), start=1):
        service_content = {
            'ten_dich_vu': f"3.{idx} {data['ten_dich_vu']}",
            'noi_dung': []
        }
        
        for noi_dung_idx, (ten_noi_dung, chi_tiet) in enumerate(data['noi_dung'], start=1):
            content_item = {
                'title': f"3.{idx}.{noi_dung_idx}" if num_services > 1 else f"3.{noi_dung_idx}",
                'ten_noi_dung': ten_noi_dung,
                'chi_tiet': []
            }
            
            if chi_tiet:
                chi_tiet_parts = chi_tiet.split('\\n')
                for detail_idx, part in enumerate(chi_tiet_parts, start=1):
                    if part.strip():
                        detail_number = f"3.{idx}.{noi_dung_idx}.{detail_idx}" if num_services > 1 else f"3.{noi_dung_idx}.{detail_idx}"
                        content_item['chi_tiet'].append({
                            'number': detail_number,
                            'text': part.strip()
                        })
            service_content['noi_dung'].append(content_item)
        
        preview_data['dieu_3'].append(service_content)
    
    return preview_data

def generate_contract_word(service_details, contract_number, sign_date):
    doc = Document()
    
    # Add header with logo and company info
    section = doc.sections[0]
    header = section.header
    header_table = header.add_table(1, 2, width=Inches(8))
    
    # Logo cell
    logo_cell = header_table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_path = os.path.join(Config.STATIC_FOLDER, 'images', 'logo_contract.png')
    if os.path.exists(logo_path):
        logo_run.add_picture(logo_path, width=Inches(2))
    else:
        print(f"Warning: Logo file not found at {logo_path}")
    
    # Company info cell
    info_cell = header_table.cell(0, 1)
    info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_cell.add_paragraph('Hanoi: 11 Dich Vong Hau, Cau Giay District').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_cell.add_paragraph('Ho Chi Minh: 105D Ben Van Don, District 4').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_cell.add_paragraph('Phone: +(84) 908-535-898').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_cell.add_paragraph('Email: univn@exumuni.com').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add Vietnamese title
    doc.add_paragraph('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Độc lập - Tự do - Hạnh phúc').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add English title
    eng_title = doc.add_paragraph('SOCIALIST REPUBLIC OF VIETNAM')
    eng_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eng_title.runs[0].italic = True
    
    eng_subtitle = doc.add_paragraph('Independence - Freedom - Happiness')
    eng_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eng_subtitle.runs[0].italic = True
    
    # Add separator
    doc.add_paragraph('———oOo———').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add contract title
    doc.add_paragraph('HỢP ĐỒNG DỊCH VỤ').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add contract number
    contract_para = doc.add_paragraph()
    contract_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contract_para.add_run(f'Số: {contract_number}').italic = True
    
    # Add English contract title
    eng_contract_title = doc.add_paragraph('SERVICE CONTRACT')
    eng_contract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eng_contract_title.runs[0].italic = True
    
    # Add English contract number
    eng_contract_para = doc.add_paragraph()
    eng_contract_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eng_contract_para.add_run(f'No. {contract_number}').italic = True
    
    # Add contract date and intro Vietnamese
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    date_text = f'Hợp đồng này được lập ngày {sign_date} tại trụ sở của CÔNG TY TNHH TƯ VẤN UNI.'
    date_para.add_run(date_text)
    
    intro_para = doc.add_paragraph()
    intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_text = 'Sau khi bàn bạc thỏa thuận, hai bên nhất trí ký Hợp đồng dịch vụ tư vấn với các điều khoản sau:'
    intro_para.add_run(intro_text)
    
    # Add English contract date and intro
    eng_date_para = doc.add_paragraph()
    eng_date_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    eng_date_text = f'This contract was made on {sign_date} at UNI CONSULTING CO., LTD.'
    eng_date_run = eng_date_para.add_run(eng_date_text)
    eng_date_run.italic = True
    
    eng_intro_para = doc.add_paragraph()
    eng_intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    eng_intro_text = 'After discussing the agreement, the two sides agreed to sign a consulting service contract with the following terms:'
    eng_intro_run = eng_intro_para.add_run(eng_intro_text)
    eng_intro_run.italic = True
    
    # Add content section
    doc.add_heading('ĐIỀU 3: NỘI DUNG CÔNG VIỆC / ARTICLE 3: CONTENT OF WORK', level=1)
    
    # Process services and their content
    num_services = len(service_details)
    for idx, (service_id, data) in enumerate(service_details.items(), start=1):
        if num_services > 1:
            # Multiple services case
            service_para = doc.add_paragraph()
            service_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            service_para.add_run(f'3.{idx} {data["ten_dich_vu"]}:').bold = True
            
            for noi_dung_idx, (ten_noi_dung, chi_tiet) in enumerate(data['noi_dung'], start=1):
                if ten_noi_dung:
                    content_para = doc.add_paragraph()
                    content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    content_para.add_run(f'3.{idx}.{noi_dung_idx} {ten_noi_dung}:').bold = True
                    
                    if chi_tiet:
                        chi_tiet_parts = chi_tiet.split('\\n')
                        for detail_idx, part in enumerate(chi_tiet_parts, start=1):
                            if part.strip():
                                detail_para = doc.add_paragraph()
                                detail_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                detail_para.add_run(f'3.{idx}.{noi_dung_idx}.{detail_idx} {part.strip()}')
        else:
            # Single service case
            for noi_dung_idx, (ten_noi_dung, chi_tiet) in enumerate(data['noi_dung'], start=1):
                if ten_noi_dung:
                    content_para = doc.add_paragraph()
                    content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    content_para.add_run(f'3.{noi_dung_idx} {ten_noi_dung}:').bold = True
                    
                    if chi_tiet:
                        chi_tiet_parts = chi_tiet.split('\\n')
                        for detail_idx, part in enumerate(chi_tiet_parts, start=1):
                            if part.strip():
                                detail_para = doc.add_paragraph()
                                detail_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                detail_para.add_run(f'3.{noi_dung_idx}.{detail_idx} {part.strip()}')

    # Save the document
    output_dir = os.path.join(Config.UPLOAD_FOLDER, 'contracts')
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, 'hop_dong.docx')
    doc.save(file_path)
    return file_path

def list_contracts():
    """
    Lấy danh sách hợp đồng từ database
    """
    try:
        query = text("""
            SELECT 
                h.id_hop_dong,
                h.ten_hop_dong,
                h.so_hop_dong,
                h.ngay_ky,
                GROUP_CONCAT(DISTINCT d.ten_dich_vu) as dich_vu
            FROM hop_dong h
            LEFT JOIN hop_dong_dich_vu hdv ON h.id_hop_dong = hdv.id_hop_dong
            LEFT JOIN dich_vu d ON hdv.id_dich_vu = d.id_dich_vu
            GROUP BY h.id_hop_dong, h.ten_hop_dong, h.so_hop_dong, h.ngay_ky
            ORDER BY h.id_hop_dong DESC
        """)
        
        result = db.session.execute(query)
        contracts = []
        for row in result:
            ngay_ky = row.ngay_ky
            if isinstance(ngay_ky, str):
                # Nếu đã là string, giữ nguyên
                formatted_date = ngay_ky
            else:
                # Nếu là date object, format lại
                formatted_date = ngay_ky.strftime('%d/%m/%Y') if ngay_ky else ''
                
            contracts.append({
                'id': row.id_hop_dong,
                'ten_hop_dong': row.ten_hop_dong,
                'so_hop_dong': row.so_hop_dong,
                'ngay_ky': formatted_date,
                'dich_vu': row.dich_vu if row.dich_vu else ''
            })
        return contracts
    except Exception as e:
        print(f"Error in list_contracts: {str(e)}")
        raise

def get_contract_details(id_hop_dong):
    """
    Lấy thông tin chi tiết của một hợp đồng theo ID
    """
    try:
        # Lấy thông tin cơ bản của hợp đồng
        contract_query = text("""
            SELECT 
                h.id_hop_dong,
                h.ten_hop_dong,
                h.so_hop_dong,
                h.ngay_ky
            FROM hop_dong h
            WHERE h.id_hop_dong = :id
        """)
        
        contract = db.session.execute(contract_query, {'id': id_hop_dong}).first()
        
        if contract is None:
            return None
            
        # Lấy danh sách dịch vụ đã chọn của hợp đồng
        selected_services_query = text("""
            SELECT 
                d.id_dich_vu
            FROM dich_vu d
            JOIN hop_dong_dich_vu hdv ON d.id_dich_vu = hdv.id_dich_vu
            WHERE hdv.id_hop_dong = :id
        """)
        
        selected_services = db.session.execute(selected_services_query, {'id': id_hop_dong}).fetchall()
        selected_service_ids = [service.id_dich_vu for service in selected_services]
            
        # Handle date formatting
        ngay_ky = contract.ngay_ky
        if isinstance(ngay_ky, str):
            # If it's already a string in dd/mm/yyyy format, keep it
            if '/' in ngay_ky:
                formatted_date = ngay_ky
            else:
                # Try to parse the string as a date
                try:
                    from datetime import datetime
                    parsed_date = datetime.strptime(ngay_ky, '%Y-%m-%d')
                    formatted_date = parsed_date.strftime('%d/%m/%Y')
                except ValueError:
                    formatted_date = ngay_ky
        else:
            # If it's a date object, format it
            formatted_date = ngay_ky.strftime('%d/%m/%Y') if ngay_ky else ''
            
        return {
            'id': contract.id_hop_dong,
            'ten_hop_dong': contract.ten_hop_dong,
            'so_hop_dong': contract.so_hop_dong,
            'ngay_ky': formatted_date,
            'id_dich_vu': selected_service_ids
        }
    except Exception as e:
        print(f"Error in get_contract_details: {str(e)}")
        raise

def upsert_contract(id_hop_dong, ten_hop_dong, so_hop_dong, ngay_ky, dich_vu_ids):
    """
    Tạo mới hoặc cập nhật hợp đồng và các dịch vụ liên quan
    """
    try:
        if id_hop_dong:
            # Cập nhật hợp đồng
            update_query = text("""
                UPDATE hop_dong
                SET ten_hop_dong = :ten_hop_dong,
                    so_hop_dong = :so_hop_dong,
                    ngay_ky = :ngay_ky
                WHERE id_hop_dong = :id_hop_dong
                RETURNING id_hop_dong
            """)
            
            result = db.session.execute(update_query, {
                'id_hop_dong': id_hop_dong,
                'ten_hop_dong': ten_hop_dong,
                'so_hop_dong': so_hop_dong,
                'ngay_ky': ngay_ky
            })
            
            if not result.fetchone():
                return {'error': 'Không tìm thấy hợp đồng để cập nhật'}
                
        else:
            # Tạo mới hợp đồng
            insert_query = text("""
                INSERT INTO hop_dong (ten_hop_dong, so_hop_dong, ngay_ky)
                VALUES (:ten_hop_dong, :so_hop_dong, :ngay_ky)
                RETURNING id_hop_dong
            """)
            
            result = db.session.execute(insert_query, {
                'ten_hop_dong': ten_hop_dong,
                'so_hop_dong': so_hop_dong,
                'ngay_ky': ngay_ky
            })
            
            id_hop_dong = result.fetchone()[0]

        # Cập nhật dịch vụ cho hợp đồng
        if dich_vu_ids is not None:
            # Xóa các liên kết cũ
            delete_old_links = text("""
                DELETE FROM hop_dong_dich_vu
                WHERE id_hop_dong = :id_hop_dong
            """)
            db.session.execute(delete_old_links, {'id_hop_dong': id_hop_dong})
            
            # Thêm các liên kết mới
            if dich_vu_ids:
                # Chuyển đổi dich_vu_ids thành list nếu là string
                if isinstance(dich_vu_ids, str):
                    dich_vu_ids = [int(dich_vu_ids)]
                elif not isinstance(dich_vu_ids, list):
                    dich_vu_ids = list(dich_vu_ids)
                
                insert_new_links = text("""
                    INSERT INTO hop_dong_dich_vu (id_hop_dong, id_dich_vu)
                    VALUES (:id_hop_dong, :id_dich_vu)
                """)
                
                for dich_vu_id in dich_vu_ids:
                    try:
                        db.session.execute(insert_new_links, {
                            'id_hop_dong': id_hop_dong,
                            'id_dich_vu': int(dich_vu_id)
                        })
                    except ValueError as e:
                        print(f"Error converting service ID: {dich_vu_id}, Error: {str(e)}")
                        continue
            
        db.session.commit()
        return {'message': 'Cập nhật hợp đồng thành công', 'id': id_hop_dong}
        
    except Exception as e:
        db.session.rollback()
        print(f"Error in upsert_contract: {str(e)}")
        return {'error': f'Lỗi khi cập nhật hợp đồng: {str(e)}'}

def delete_contract(id_hop_dong):
    """
    Xóa hợp đồng theo id
    """
    try:
        # Xóa các liên kết trong bảng trung gian
        delete_links = text("""
            DELETE FROM hop_dong_dich_vu
            WHERE id_hop_dong = :id_hop_dong
        """)
        db.session.execute(delete_links, {'id_hop_dong': id_hop_dong})
        
        # Xóa hợp đồng
        delete_query = text("""
            DELETE FROM hop_dong 
            WHERE id_hop_dong = :id_hop_dong
            RETURNING id_hop_dong
        """)
        
        result = db.session.execute(delete_query, {'id_hop_dong': id_hop_dong})
        deleted = result.fetchone()
        
        if deleted:
            db.session.commit()
            return {'message': 'Xóa hợp đồng thành công'}
        else:
            return {'error': 'Không tìm thấy hợp đồng để xóa'}
            
    except Exception as e:
        db.session.rollback()
        return {'error': f'Lỗi khi xóa hợp đồng: {str(e)}'}

def get_available_services(current_contract_id=None):
    """
    Lấy danh sách dịch vụ chưa được gán cho hợp đồng nào
    """
    try:
        if current_contract_id:
            # Nếu đang sửa hợp đồng, lấy các dịch vụ chưa được gán và dịch vụ của hợp đồng hiện tại
            query = text("""
                SELECT id_dich_vu, ten_dich_vu
                FROM dich_vu
                WHERE id_hop_dong IS NULL
                OR id_hop_dong = :contract_id
            """)
            params = {'contract_id': current_contract_id}
        else:
            # Nếu đang tạo mới, chỉ lấy các dịch vụ chưa được gán
            query = text("""
                SELECT id_dich_vu, ten_dich_vu
                FROM dich_vu
                WHERE id_hop_dong IS NULL
            """)
            params = {}
        
        services = db.session.execute(query, params)
        return [
            {'id': service.id_dich_vu, 'name': service.ten_dich_vu}
            for service in services
        ]
    except Exception as e:
        print(f"Error in get_available_services: {str(e)}")
        raise
